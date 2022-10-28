from docx.api import Document
import pandas as pd

TOTAL = 0

df = pd.DataFrame()
document = Document("Week 8.docx")
tables = document.tables

def read_docx(document):
    total_rows = list()
    for table in document.tables:
        for row in table.rows:
            text = [" ".join(cell.text.strip().split("\n")) for cell in row.cells if cell.text != ""] # if cell.text not in ("BYE", "")
            total_rows.append(text)
    return total_rows


def split_headers(document):
    buffer = list()
    for i in range(len(document[0])):
        current = document[0][i].strip()
        if i == 0:
            TOTAL = current.split("$")[1]
            current = "Name"
        buffer.append(" ".join(current.split("\n")))
    # if a bye column is in list it will be removed
    headers = [item for item in buffer if "BYE" not in item]
    document[0] = headers
    return headers
    
rows = read_docx(document)
head = split_headers(rows)


df = pd.DataFrame(rows[1:], columns=head)
df.to_csv("test.csv", index=False)



